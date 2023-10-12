import os, sys
import tkinter as tk
from tkinter import messagebox
from tkinter import filedialog
from docx import Document
from docx.shared import Mm, Pt
from datetime import datetime
from src.config import *
from src.exception import CustomException
from src.logger import logging
from src.utils import *
from src.config import *
from src.components import imgIngestion, imgFrame


# Function to check if an image exceeds the size limit
def is_exceed_limit(image_path:str) -> bool:
    """
    This helps converting process avoid large resolution/size images
    """
    try:
        _ = Image.open(image_path)
        return True
    except Exception as e:
        return False


# Function to confirm with the user whether to continue with images exceeding the size limit
def confirm_large_images(exceeding_images:list) -> bool:
    """
    This function araise when there is an image exceeding limit size found.
    Give user options whether to continue with passing those images or stop the converting process.
    """
    try:
        if not exceeding_images:
            return True  # No images exceeded the size limit, continue with the process

        # Generate a message with the list of exceeding images
        message = f"Found images exceed the size limit and will be excluded.\n\n"
        message += "\n\nDo you want to continue?"

        response = messagebox.askquestion("Image Size Warning", message)
        return response.lower() == "yes"

    except Exception as e:
        raise CustomException(e, sys)


class Converter:
    def __init__(self) -> None:
        self.output_folder:str 
        self.copies:int
        
    
    # Define a function to handle the Convert button click
    def convert_button_click(self, inputImages, output_folder_var, copies_var):
        try:
            self.output_folder = output_folder_var.get()
            self.copies = int(copies_var.get())

            if not self.output_folder:
                messagebox.showwarning("Warning", f"Please, select the output folder!")
                return
            
            if not os.path.exists(self.output_folder):
                os.makedirs(self.output_folder)

            if not inputImages['path']:
                messagebox.showwarning("Warning", f"No image has been selected!")
                return

            self.create_word_document(inputImages)

        except Exception as e:
            raise CustomException(e, sys)


    def create_word_document(self, inputImages:dict) -> None:
        try:
            doc = Document()
            exceeding_images = []

            loggingInfo(f"Converting {len(inputImages['path'])} images in process.")
            
            # Docx set up
            section = doc.sections[0]
            section.left_margin = Mm(12.7)  
            section.right_margin = Mm(12.7)  
            section.top_margin = Mm(12.7)  
            section.bottom_margin = Mm(12.7)  

            style = doc.styles['Normal']
            style.paragraph_format.space_before = Pt(0)  
            style.paragraph_format.space_after = Pt(3)  
            
            current_time = datetime.now().strftime('%Y-%m-%d_%H-%M-%S')
            output_path = os.path.join(self.output_folder, f"{current_time}.docx")

            
            for image_path, copies, shape in zip(inputImages['path'], inputImages['copies'], inputImages['shape']):
                copies = int(copies)
                is_exceed = is_exceed_limit(image_path)
                
                if not is_exceed and len(exceeding_images)!=0:
                    exceeding_images.append(image_path)    
                    continue
                
                if not is_exceed:         
                    exceeding_images.append(image_path)   
                    stop_sign = confirm_large_images(exceeding_images)

                    if not stop_sign:
                        messagebox.showinfo("Info", "Operation canceled.")
                        loggingInfo("User cancaled converting process due to high resolution images.")
                        return
                
                    continue
                
                # Crop the image to square and convert to RGB
                img = Image.open(image_path)

                iwidth, iheight = img.size
                xc, yc = iwidth//2, iheight//2

                ingestor = imgIngestion.ImageIngestion()
                ingestor.initiate_image_obj(img, px_height=min(1000, iheight), xc=xc, yc=yc)

                img = ingestor.crop_by_shape(shape=shape)
                
                if img is None:
                    continue  # Skip this image and continue with the next one

                # Add the picture to the paragraph with copies           
                img.save("temp_image.jpg")  # Save the square image temporarily

                # Repeat the image in the same line
                run = doc.add_paragraph().add_run()
                for _ in range(copies):
                    run.add_picture("temp_image.jpg", width=Mm(16), height=Mm(16))
                    run.add_text(" ")  
                os.remove("temp_image.jpg")  # Remove the temporary file
            
            if len(exceeding_images)>0:
                loggingInfo(f"Found {len(exceeding_images)} high resolution images")
            
            if len(exceeding_images) == len(inputImages['path']):
                messagebox.showwarning("Warning", "All images are exceeding the maximum size. \
                                    Please, choose another folder or resize your images!")
                loggingInfo("Converting process stopped. All images are high resolution.")
            else:    
                doc.save(output_path)
                messagebox.showinfo("Success", f"Document saved to:\n{output_path}")
                loggingInfo(f"Converted {len(inputImages) - len(exceeding_images)} images completely. Document saved to: {output_path}")
        
        except Exception as e:
            raise CustomException(e, sys)

    
            
    