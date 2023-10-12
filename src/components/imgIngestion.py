import os, sys
from PIL import Image
from tkinter import messagebox
from src.logger import logging
from src.exception import CustomException
from src.utils import *
from src.config import *
from docx import Document
from docx.shared import Mm, Pt
from docx.enum.text import WD_LINE_SPACING
from datetime import datetime

def crop_img(img, px_width, px_height, xc, yc) -> Image.Image:
        try:
            # Calculate the crop box coordinates
            left = xc - (px_width / 2)
            top = yc - (px_height / 2)
            right = xc + (px_width / 2)
            bottom = yc + (px_height / 2)
            # Crop the image
            cropped_img = img.crop((left, top, right, bottom))
            cropped_img = cropped_img.convert("RGB")

            return cropped_img
        except Exception as e:
            raise CustomException(e, sys)
        

class ImageIngestion:
    def __init__(self) -> None:
        self.img:Image.Image
        self.px_width:int
        self.px_height:int
        self.xc:int
        self.yc:int


    def initiate_image_obj(self, img:Image.Image,
                  px_height:int, xc:int, yc:int) -> None:
        try:
            self.img = img
            self.px_height = px_height
            self.xc = xc
            self.yc = yc

            loggingInfo(f"[ImageIngestion] Initiate image with pixel height [{px_height}], xc [{xc}] and yc [{yc}] completely.")
        
        except Exception as e:
            raise CustomException(e, sys)
        
    
    def crop_combo(self):
        try:
            crop_list = {}
            loggingInfo("[ImageIngestion] Cropping Combo in process.")
            x_ori, y_ori = self.xc - self.px_width//2, self.yc - self.px_height//2
            height = self.px_height//4        
            for key in  COMBO_GRID.keys():
                y = y_ori + int(height*(COMBO_GRID[key][0]+1/2))
                width = SHAPE_RATIOS[key] * height
                x = x_ori + int(width*(COMBO_GRID[key][1]+1/2))
                loggingInfo(f"[ImageIngestion] Cropping Combo key {key}, width = {width}, height = {height}, x = {x}, y = {y}.")
                cropped_img = crop_img(self.img, width, height, x, y)
                # cropped_img.show()
                crop_list[key] = cropped_img

            return crop_list
        except Exception as e:
            raise CustomException(e, sys)



    def crop_by_shape(self, shape:str) -> Image.Image:
        try:
            if shape not in SHAPE_RATIOS.keys():
                loggingInfo(f"[ImageIngestion] Undefined shape input: {shape}.")
                raise AssertionError(f"[ImageIngestion] Shape must be either one of these: {SHAPE_RATIOS.keys()}")
            self.px_width = SHAPE_RATIOS[shape] * self.px_height
            
            if shape != 'Combo':
                cropped_img = crop_img(self.img, self.px_width, self.px_height, self.xc, self.yc)
                loggingInfo(f"[ImageIngestion] Crop image by shape of {shape} completely.")
                return cropped_img
            else:
                cropped_imgs = self.crop_combo()
                loggingInfo(f"[ImageIngestion] Crop image by Combo completely.")
                return cropped_imgs
            
            
        
        except Exception as e:
            raise CustomException(e, sys)


## TEST     
if __name__=="__main__":
    paths = ["..\\img\\sample1.jpg", "..\\img\\sample2.jpg"]
    doc = Document()
        
    # Set page margins (0.5 inch for all borders)
    section = doc.sections[0]
    section.left_margin = Mm(12.7)  # 0.5 inch
    section.right_margin = Mm(12.7)  # 0.5 inch
    section.top_margin = Mm(12.7)  # 0.5 inch
    section.bottom_margin = Mm(12.7)  # 0.5 inch
    
    style = doc.styles['Normal']
    style.paragraph_format.space_before = Pt(0)  # Set space before the paragraph to 3 points
    style.paragraph_format.space_after = Pt(3)  # Set space after the paragraph to 3 points


    current_time = datetime.now().strftime('%Y-%m-%d_%H-%M-%S')
    output_path = os.path.join("C:\\Users\\USER\\Desktop", f"{current_time}.docx")

    for path in paths:
        img = Image.open(path)

        iwidth, iheight = img.size

        print(iwidth, iheight)
        xc, yc = iwidth//2, iheight//2

        tool = ImageIngestion()
        
        tool.initiate_image_obj(img, px_height=min(1000, iheight), xc=xc, yc=yc)

        crop_imgs = tool.crop_by_shape(shape='Combo')
        

        for key, img in crop_imgs.items():
            img.save("temp_square_image.jpg")  # Save the square image temporarily
            
            if key != 'Win':
                paragraph = doc.add_paragraph()
                run = paragraph.add_run()
                # Set line spacing for the paragraph
                paragraph_format = paragraph.paragraph_format
                paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

            run.add_picture("temp_square_image.jpg", height=Mm(IMG_SHAPES[key][0]))        
            run.add_text(" ")
            
            os.remove("temp_square_image.jpg")  # Remove the temporary file
        
        
    doc.save(output_path)
    messagebox.showinfo("Success", f"Document saved completely")

    
