import os, sys
import tkinter as tk
from tkinter import messagebox
from tkinter import filedialog
from docx import Document
from docx.shared import Mm
from datetime import datetime
from src.config import *
from src.exception import CustomException
from src.logger import logging
from src.utils import *
from src.config import IMG_EXTENSIONS


class Resize:
    def __init__(self) -> None:
        self.output_folder:str = None
        self.input_folder:str = None
        self.repetition_lines:int = None
        self.images_list:list = []

    def add_images_button(self):
        
        return 
        
    # Define a function to handle the Convert button click
    def convert_button_click(self, input_folder_var, output_folder_var, repetition_images):
        self.input_folder = input_folder_var.get()
        self.output_folder = output_folder_var.get()
        repetition_lines_str = repetition_images.get()
        loggingInfo("User chose parameters")

        if not self.input_folder:
            messagebox.showwarning("Warning", f"Please, select the input folder!")
            return

        if not repetition_lines_str:
            messagebox.showwarning("Warning", "Please enter a value for Repetition Lines.")
            return

        try:
            self.repetition_lines = int(repetition_lines_str)  
        except Exception as e:
            messagebox.showwarning("Warning", "Repetition Lines must be a valid integer.")
            CustomException(e, sys)
            return

        self.create_word_document()

    def create_word_document(self) -> None:
        doc = Document()
        exceeding_images = []
        
        if not os.path.exists(self.output_folder):
            os.makedirs(self.output_folder)
        
        image_files = []
        for filename in os.listdir(self.input_folder):
            if filename.lower().endswith(tuple(IMG_EXTENSIONS)):
                image_files.append(os.path.join(self.input_folder, filename))
        
        if not image_files:
            messages = f"No image files found in the input folder!"
            messagebox.showwarning("Warning", messages)
            loggingInfo(messages)
            return

        loggingInfo(f"Load {len(image_files)} images completely from {self.output_folder}")
        
        # Set page margins (0.5 inch for all borders)
        section = doc.sections[0]
        section.left_margin = Mm(12.7)  # 0.5 inch
        section.right_margin = Mm(12.7)  # 0.5 inch
        section.top_margin = Mm(12.7)  # 0.5 inch
        section.bottom_margin = Mm(12.7)  # 0.5 inch
        
        current_time = datetime.now().strftime('%Y-%m-%d_%H-%M-%S')
        output_path = os.path.join(self.output_folder, f"{current_time}.docx")

        
        for image_file in image_files:
            
            is_exceed = is_exceed_limit(image_file)
            
            if not is_exceed and len(exceeding_images)!=0:
                exceeding_images.append(image_file)    
                continue
            
            if not is_exceed:         
                exceeding_images.append(image_file)   
                stop_sign = confirm_large_images(exceeding_images)

                if not stop_sign:
                    messagebox.showinfo("Info", "Operation canceled.")
                    return
            
                continue
            
            # Crop the image to square and convert to RGB
            img = crop_to_square(image_file)
            
            if img is None:
                continue  # Skip this image and continue with the next one

            # Add the picture to the paragraph with repetitions
            img.save("temp_square_image.jpg")  # Save the square image temporarily
            
            # Repeat the image in the same line
            run = doc.add_paragraph().add_run()
            for _ in range(self.repetition_lines):
                run.add_picture("temp_square_image.jpg", width=Mm(16), height=Mm(16))
                run.add_text(" ")
            
            
            os.remove("temp_square_image.jpg")  # Remove the temporary file
        
        
        if len(exceeding_images) == len(image_files):
            messagebox.showwarning("Warning", "All images are exceeding the maximum size. \
                                   Please, choose another folder or resize your images!")
            loggingInfo("Folder contains high resolution images")
            
        doc.save(output_path)
        messagebox.showinfo("Success", f"Document saved to:\n{output_path}")
        

    
            
    